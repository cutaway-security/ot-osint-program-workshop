"""Generate monitoring-checklist-template.docx with checkboxes and tables."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Style configuration
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    hs.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark header background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A365D",
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)

    # Data rows
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_checkbox_item(doc, text, level=0):
    """Add a checkbox item using Unicode ballot box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27 * level) if level else None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Unicode ballot box
    run = p.add_run("\u2610 ")
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    return p


# ============================================================
# Title
# ============================================================
title = doc.add_heading("Pull-Based Monitoring Checklist", level=1)

p = doc.add_paragraph()
p.add_run("Instructions: ").bold = True
p.add_run("Copy this template and customize the queries, baseline references, and finding criteria for your organization. Use the checklist format during each monitoring cycle to ensure consistent execution.")

doc.add_paragraph()

# ============================================================
# Weekly Monitoring Checklist
# ============================================================
doc.add_heading("Weekly Monitoring Checklist", level=2)

p = doc.add_paragraph()
run = p.add_run("Target time: ")
run.bold = True
p.add_run("Under 30 minutes  |  ")
run2 = p.add_run("Schedule: ")
run2.bold = True
p.add_run("Every [day]  |  ")
run3 = p.add_run("Owner: ")
run3.bold = True
p.add_run("[name]")

p2 = doc.add_paragraph()
p2.add_run("Checks ordered by OT relevance -- remote access and edge device checks first.").italic = True

# --- Check 1 ---
doc.add_heading("Check 1: Shodan/Censys Remote Access Review", level=3)

add_checkbox_item(doc, "Run query in Shodan / Censys")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("Shodan (shodan.io) / Censys (search.censys.io)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("[your exact query -- e.g., net:x.x.x.x/24 port:443,8443,10443]")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Section 1, rows [X-Y] (remote access services)")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New Findings", "Classification", "Action Taken"],
    [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
    [3, 6, 4, 6],
)

p = doc.add_paragraph()
run = p.add_run("OT-relevant finding: ")
run.bold = True
p.add_run("Any change to VPN login page, new port on firewall management IP, version change on edge device. Version downgrade or unexpected service warrants immediate investigation.")

p = doc.add_paragraph()
run = p.add_run("Expected noise: ")
run.bold = True
p.add_run("Minor HTTP header changes, certificate renewals on same service.")

# --- Check 2 ---
doc.add_heading("Check 2: CISA KEV Review", level=3)

add_checkbox_item(doc, "Review CISA KEV additions from past 7 days")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("CISA KEV Catalog (cisa.gov/known-exploited-vulnerabilities-catalog)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("Review additions from past 7 days. Cross-reference against asset inventory.")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Module 4 vulnerability correlation table")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New KEV Entries", "Match to Inventory", "Priority", "Action Taken"],
    [["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""]],
    [3, 4, 4, 3, 5],
)

p = doc.add_paragraph()
run = p.add_run("OT-relevant finding: ")
run.bold = True
p.add_run("Any KEV entry matching a product/vendor in your asset inventory. P0 if the asset is internet-exposed.")

# --- Check 3 ---
doc.add_heading("Check 3: Certificate Transparency", level=3)

add_checkbox_item(doc, "Search crt.sh for new certificates")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("crt.sh")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("%.[your domain]")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Section 1, subdomain list")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New Subdomains", "Classification", "Action Taken"],
    [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
    [3, 6, 4, 6],
)

p = doc.add_paragraph()
run = p.add_run("OT-relevant finding: ")
run.bold = True
p.add_run("New subdomains suggesting remote access (vpn*, ras*, remote*), new OT-adjacent applications, or shadow IT.")

# --- Check 4 ---
doc.add_heading("Check 4: Breach Database", level=3)

add_checkbox_item(doc, "Check HIBP for domain or Tier 1 emails")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("HaveIBeenPwned (haveibeenpwned.com)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("Domain search for [your domain] or spot-check Tier 1 emails: [list]")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Personnel Affected", "Breach Name", "Data Exposed", "Priority", "Action Taken"],
    [["", "", "", "", "", ""], ["", "", "", "", "", ""], ["", "", "", "", "", ""]],
    [3, 3.5, 3.5, 3, 2, 4],
)

# --- Check 5 ---
doc.add_heading("Check 5: Alert Backlog Processing", level=3)

add_checkbox_item(doc, "Process deferred items from daily triage")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("Alert aggregation point ([email folder / Slack / shared inbox])")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Items in Backlog", "Items Closed", "Items Escalated"],
    [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
    [3, 5, 5, 5],
)

p = doc.add_paragraph()
run = p.add_run("Rule: ")
run.bold = True
p.add_run("No item remains in backlog longer than one week.")

# --- Check 6 ---
doc.add_heading("Check 6: Baseline Update", level=3)

add_checkbox_item(doc, "Update baseline document if any changes found")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Section Updated", "Change Description"],
    [["", "", ""], ["", "", ""], ["", "", ""]],
    [3, 5, 10],
)

# ============================================================
# Monthly Monitoring Checklist
# ============================================================
doc.add_heading("Monthly Monitoring Checklist", level=2)

p = doc.add_paragraph()
run = p.add_run("Target time: ")
run.bold = True
p.add_run("60-90 minutes  |  ")
run2 = p.add_run("Schedule: ")
run2.bold = True
p.add_run("[first Monday / last Friday]  |  ")
run3 = p.add_run("Owner: ")
run3.bold = True
p.add_run("[name]")

# --- Monthly Check 1 ---
doc.add_heading("Check 1: Full Subdomain Re-enumeration (20 min)", level=3)

add_checkbox_item(doc, "Run full subdomain enumeration on all domains")

p = doc.add_paragraph()
run = p.add_run("Tools: ")
run.bold = True
p.add_run("crt.sh, DNSDumpster, SecurityTrails")

p = doc.add_paragraph()
run = p.add_run("Domains: ")
run.bold = True
p.add_run("[list all domains]")

add_table(doc,
    ["Date", "New Subdomains", "Removed Subdomains", "Service Changes", "Action Taken"],
    [["", "", "", "", ""], ["", "", "", "", ""]],
    [3, 4, 4, 4, 4],
)

# --- Monthly Check 2 ---
doc.add_heading("Check 2: Deep Shodan/Censys Scan (15 min)", level=3)

add_checkbox_item(doc, "Run broader Shodan/Censys queries")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("[broader queries -- org name, ASN, netblocks]")

p = doc.add_paragraph()
run = p.add_run("Focus: ")
run.bold = True
p.add_run("New management interfaces and remote access on OT boundary.")

add_table(doc,
    ["Date", "New Services", "Changed Services", "Action Taken"],
    [["", "", "", ""], ["", "", "", ""]],
    [3, 5, 5, 5],
)

# --- Monthly Check 3 ---
doc.add_heading("Check 3: Vulnerability Re-correlation (20 min)", level=3)

add_checkbox_item(doc, "Re-run vulnerability queries for all inventoried products")

p = doc.add_paragraph()
run = p.add_run("Tools: ")
run.bold = True
p.add_run("NVD, CISA KEV, vendor PSIRTs, ICS Advisory Project")

p = doc.add_paragraph()
run = p.add_run("Products: ")
run.bold = True
p.add_run("[list from Module 4 asset inventory]")

add_table(doc,
    ["Date", "New CVEs", "KEV Changes", "Priority Updates", "Action Taken"],
    [["", "", "", "", ""], ["", "", "", "", ""]],
    [3, 4, 4, 4, 4],
)

# --- Monthly Check 4 ---
doc.add_heading("Check 4: Personnel Exposure Refresh (15 min)", level=3)

add_checkbox_item(doc, "Re-check Tier 1/2 personnel against breach databases")
add_checkbox_item(doc, "Review personnel changes (new hires, departures, role changes)")

add_table(doc,
    ["Date", "New Breaches", "Personnel Changes", "Action Taken"],
    [["", "", "", ""], ["", "", "", ""]],
    [3, 5, 5, 5],
)

# --- Monthly Check 5 ---
doc.add_heading("Check 5: Alert Configuration Review (10 min)", level=3)

add_checkbox_item(doc, "Verify all push alerts are active and delivering")
add_checkbox_item(doc, "Review query relevance")

add_table(doc,
    ["Date", "Alerts Verified", "Changes Made"],
    [["", "", ""], ["", "", ""]],
    [3, 6, 8],
)

# ============================================================
# Cycle Summary
# ============================================================
doc.add_heading("Cycle Summary", level=2)

p = doc.add_paragraph("Complete this section at the end of each monitoring cycle.")

add_table(doc,
    ["Cycle Date", "Type", "Total Findings", "P0/P1", "Baseline Updates", "Time Spent", "Analyst"],
    [
        ["", "Weekly", "", "", "", "", ""],
        ["", "Weekly", "", "", "", "", ""],
        ["", "Monthly", "", "", "", "", ""],
        ["", "Monthly", "", "", "", "", ""],
    ],
    [3, 2.5, 3, 2, 3, 2.5, 3],
)

doc.save("/home/cutaway/Projects/ot-osint-program-workshop/templates/downloads/monitoring-checklist-template.docx")
print("monitoring-checklist-template.docx created")
