"""Generate runbook-template.docx with [CORE]/[FULL] labels and checkboxes."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Style configuration
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)

for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    hs.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A365D", qn("w:val"): "clear",
        })
        shading.append(shading_elm)
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_step(doc, label, text, indent=False):
    """Add a checkbox step with [CORE] or [FULL] label."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_box = p.add_run("\u2610 ")
    run_box.font.size = Pt(11)
    run_label = p.add_run(f"[{label}] ")
    run_label.font.bold = True
    run_label.font.size = Pt(11)
    if label == "CORE":
        run_label.font.color.rgb = RGBColor(0x05, 0x96, 0x69)  # Green
    else:
        run_label.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)  # Blue
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    return p


def add_detail(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(value)
    run2.font.size = Pt(10)
    return p


# ============================================================
# Title
# ============================================================
doc.add_heading("Operational Runbook", level=1)

p = doc.add_paragraph()
p.add_run("Instructions: ").bold = True
p.add_run("Copy this template and customize with your organization-specific details. Replace all [bracketed items] with actual tool queries, personnel names, and baseline references. Mark each procedure as [CORE] or [FULL].")

doc.add_paragraph()

p = doc.add_paragraph()
run_core = p.add_run("[CORE]")
run_core.bold = True
run_core.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
p.add_run(" -- Runs no matter what. Part of your minimum viable monitoring set.")

p = doc.add_paragraph()
run_full = p.add_run("[FULL]")
run_full.bold = True
run_full.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
p.add_run(" -- Runs during normal operations. Can be deferred when time is constrained.")

doc.add_paragraph()

# ============================================================
# Program Information
# ============================================================
doc.add_heading("Program Information", level=2)

info_fields = [
    ("Organization", "[name]"),
    ("Program owner", "[name, role]"),
    ("Backup owner", "[name, role]"),
    ("Escalation contact (P0/P1)", "[name, phone, email]"),
    ("Baseline document location", "[path or URL]"),
    ("Last runbook review", "[date]"),
]
for label, value in info_fields:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

doc.add_paragraph()

# ============================================================
# Daily Alert Triage
# ============================================================
doc.add_heading("Daily Alert Triage (5-10 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("[name]  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("[name]")

add_step(doc, "CORE", "Open alert aggregation point ([email folder / Slack channel / shared inbox])")
add_step(doc, "CORE", "Review new Google Alerts for: [list your configured alert queries]")
add_step(doc, "CORE", "Review CISA ICS Advisory emails -- check for advisories affecting [list your vendors]")
add_step(doc, "CORE", "Triage each alert:")
add_detail(doc, "Known baseline item", "Mark as reviewed, no action")
add_detail(doc, "New finding, low priority", "Add to weekly review queue with one-line note")
add_detail(doc, "New finding, elevated priority", "Investigate immediately using P0-P3 framework")
add_detail(doc, "New finding, critical (P0)", "Escalate to [escalation contact] immediately")
add_step(doc, "CORE", "Log triage decisions: date, source, summary, disposition")

doc.add_paragraph()

# ============================================================
# Weekly Active Monitoring
# ============================================================
doc.add_heading("Weekly Active Monitoring (25-30 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("[name]  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("[name]  |  ")
run3 = p.add_run("Day: ")
run3.bold = True
p.add_run("[e.g., Monday]")

p = doc.add_paragraph("Checks ordered by OT relevance:")
p.italic = True

add_step(doc, "CORE", "Shodan/Censys remote access check (5 min)")
add_detail(doc, "Query", "[exact Shodan/Censys query for your IPs]")
add_detail(doc, "Compare to", "Baseline Section 1, rows [X-Y] (remote access services)")
add_detail(doc, "Escalate if", "New ports, service changes, or version changes on remote access IPs")

add_step(doc, "CORE", "CISA KEV review (5 min)")
add_detail(doc, "Query", "Review CISA KEV catalog additions from past 7 days")
add_detail(doc, "Compare to", "Module 4 asset inventory -- [list your products/vendors]")
add_detail(doc, "Escalate if", "Any KEV match on an internet-exposed asset (P0)")

add_step(doc, "FULL", "Certificate transparency check (5 min)")
add_detail(doc, "Query", "crt.sh for [your domains]")
add_detail(doc, "Compare to", "Baseline Section 1, subdomain list")
add_detail(doc, "Escalate if", "New subdomains suggesting remote access (vpn*, ras*, remote*)")

add_step(doc, "FULL", "Breach database check (5 min)")
add_detail(doc, "Query", "HIBP for [domain] or spot-check Tier 1 emails: [list]")
add_detail(doc, "Compare to", "Baseline Section 2, breach status column")
add_detail(doc, "Escalate if", "Tier 1 or Tier 2 personnel in new breach with password/hash exposure")

add_step(doc, "FULL", "Alert backlog processing (5 min)")
add_detail(doc, "Action", "Review items deferred during daily triage")
add_detail(doc, "Rule", "No item should remain in backlog longer than one week")

add_step(doc, "FULL", "Baseline update (5 min)")
add_detail(doc, "Action", "Update baseline for any changes discovered this week")
add_detail(doc, "Log", "Record date and nature of each change")

doc.add_paragraph()

# ============================================================
# Monthly Full Review
# ============================================================
doc.add_heading("Monthly Full Review (60-90 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("[name]  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("[name]  |  ")
run3 = p.add_run("Week: ")
run3.bold = True
p.add_run("[e.g., first Monday]")

add_step(doc, "FULL", "Attack surface re-enumeration (20 min)")
add_detail(doc, "Tools", "crt.sh, DNSDumpster, SecurityTrails for [your domains]")
add_detail(doc, "Action", "Compare complete results to baseline. Investigate new subdomains. Remove decommissioned entries")

add_step(doc, "FULL", "Deep Shodan/Censys scan (15 min)")
add_detail(doc, "Queries", "Organization name, ASN ranges, netblock searches")
add_detail(doc, "Focus", "New management interfaces and remote access systems")

add_step(doc, "CORE", "Vulnerability re-correlation (20 min)")
add_detail(doc, "Products", "[list your edge devices and OT-adjacent products]")
add_detail(doc, "Sources", "NVD, CISA KEV, vendor PSIRTs, ICS Advisory Project")
add_detail(doc, "Action", "Update correlation table. Re-evaluate priority ratings")

add_step(doc, "FULL", "Personnel exposure refresh (15 min)")
add_detail(doc, "Action", "Re-check Tier 1 and Tier 2 against breach databases")
add_detail(doc, "Review", "New hires, departures, role changes. Add new Tier 1 personnel immediately")

add_step(doc, "FULL", "Alert configuration review (10 min)")
add_detail(doc, "Verify", "All push-based alerts still active and delivering")
add_detail(doc, "Update", "Add queries for organizational changes, new domains, new vendors")

doc.add_paragraph()

# ============================================================
# Quarterly Program Assessment
# ============================================================
doc.add_heading("Quarterly Program Assessment (2-3 hours)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("[name]  |  ")
run2 = p.add_run("Quarter: ")
run2.bold = True
p.add_run("[Q1/Q2/Q3/Q4]")

add_step(doc, "FULL", "Full baseline refresh (60 min)")
add_detail(doc, "Action", "Complete re-run of Modules 2-4 as if starting fresh")
add_detail(doc, "Compare", "Comprehensive results to current baseline. Reclassify all entries")

add_step(doc, "FULL", "Program metrics (30 min)")
add_detail(doc, "Total alerts received vs. actionable findings", "___")
add_detail(doc, "Average triage response time", "___")
add_detail(doc, "Baseline changes this quarter", "___")
add_detail(doc, "Open remediation items and aging", "___")

add_step(doc, "FULL", "Stakeholder briefing (30 min)")
add_detail(doc, "Content", "Key findings, remediation progress, emerging risks, resource requests")

add_step(doc, "FULL", "Runbook review (30 min)")
add_detail(doc, "Check", "Are procedures still accurate? Have tools changed?")
add_detail(doc, "Check", "Are there new data sources to incorporate?")
add_detail(doc, "Action", "Update cadences, queries, and contacts as needed")
add_detail(doc, "Action", "Update this runbook's 'Last runbook review' date")

doc.add_paragraph()

# ============================================================
# Change Log
# ============================================================
doc.add_heading("Change Log", level=2)

add_table(doc,
    ["Date", "Change", "Changed By"],
    [["[today]", "Initial runbook created", "[name]"], ["", "", ""], ["", "", ""]],
    [3, 10, 4],
)

doc.save("/home/cutaway/Projects/ot-osint-program-workshop/templates/downloads/runbook-template.docx")
print("runbook-template.docx created")
