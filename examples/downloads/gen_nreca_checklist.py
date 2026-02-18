"""Generate nreca-monitoring-checklist.docx pre-filled with NRECA example data."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)

for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    hs.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A365D", qn("w:val"): "clear",
        })
        shading.append(shading_elm)
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_checkbox_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("\u2610 ")
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    return p


# ============================================================
# Title
# ============================================================
doc.add_heading("Pull-Based Monitoring Checklist -- NRECA Example", level=1)

p = doc.add_paragraph()
p.add_run("NOTE: ").bold = True
p.add_run("This is a worked example using NRECA (National Rural Electric Cooperative Association) as the target organization. All data is derived from publicly available sources. Breach results are hypothetical examples created for the workshop.")

doc.add_paragraph()

# ============================================================
# Weekly Monitoring Checklist
# ============================================================
doc.add_heading("Weekly Monitoring Checklist", level=2)

p = doc.add_paragraph()
run = p.add_run("Target time: ")
run.bold = True
p.add_run("Under 30 minutes  |  ")
run2 = p.add_run("Schedule: ")
run2.bold = True
p.add_run("Every Monday  |  ")
run3 = p.add_run("Owner: ")
run3.bold = True
p.add_run("[assigned analyst]")

p2 = doc.add_paragraph()
p2.add_run("Checks ordered by OT relevance -- remote access and edge device checks first.").italic = True

# --- Check 1 ---
doc.add_heading("Check 1: Shodan/Censys Remote Access Review", level=3)

add_checkbox_item(doc, "Run query in Shodan / Censys")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("Shodan (shodan.io) / Censys (search.censys.io)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run('product:"FortiGate" port:443 "FortiOS" net:[NRECA IP ranges]')

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Section 1 -- FortiGate SSL-VPN on port 443, FortiOS 7.4.6")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New Findings", "Classification", "Action Taken"],
    [
        ["2026-02-17", "FortiOS version changed from 7.4.6 to 7.4.9", "Known-Good", "Confirmed planned patch for CVE-2025-59718. Updated baseline."],
        ["2026-02-24", "No changes detected", "--", "No action required."],
        ["", "", "", ""],
    ],
    [3, 6, 4, 6],
)

p = doc.add_paragraph()
run = p.add_run("OT-relevant finding: ")
run.bold = True
p.add_run("Any change to FortiGate VPN login page, new port on firewall management IP, FortiOS version change. Version downgrade warrants immediate investigation.")

# --- Check 2 ---
doc.add_heading("Check 2: CISA KEV Review", level=3)

add_checkbox_item(doc, "Review CISA KEV additions from past 7 days")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("CISA KEV Catalog (cisa.gov/known-exploited-vulnerabilities-catalog)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("Review additions from past 7 days. Cross-reference against: Fortinet FortiOS, Okta, PingFederate, SharePoint")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Module 4 vulnerability correlation table -- FortiGate, identity infrastructure")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New KEV Entries", "Match to Inventory", "Priority", "Action Taken"],
    [
        ["2026-02-17", "CVE-2025-59718 (FortiOS SSO bypass)", "Yes -- FortiGate", "P0", "Immediate escalation. Patched to 7.4.9 same day."],
        ["2026-02-24", "2 new KEV entries (unrelated vendors)", "No", "--", "No action required."],
        ["", "", "", "", ""],
    ],
    [3, 4, 4, 2, 5],
)

# --- Check 3 ---
doc.add_heading("Check 3: Certificate Transparency", level=3)

add_checkbox_item(doc, "Search crt.sh for new certificates")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("crt.sh")

p = doc.add_paragraph()
run = p.add_run("Queries: ")
run.bold = True
p.add_run("%.electric.coop, %.cooperative.com, %.nreca.coop, %.nrecainternational.coop")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Section 1 -- subdomain list (electric.coop: 5 subdomains, cooperative.com: 100+ subdomains)")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "New Subdomains", "Classification", "Action Taken"],
    [
        ["2026-02-17", "No new subdomains", "--", "No action required."],
        ["2026-02-24", "newportal.cooperative.com (Let's Encrypt cert)", "Needs Investigation", "New subdomain not in baseline. Investigating purpose and ownership."],
        ["", "", "", ""],
    ],
    [3, 6, 4, 6],
)

# --- Check 4 ---
doc.add_heading("Check 4: Breach Database", level=3)

add_checkbox_item(doc, "Check HIBP for domain or Tier 1 emails")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("HaveIBeenPwned (haveibeenpwned.com)")

p = doc.add_paragraph()
run = p.add_run("Query: ")
run.bold = True
p.add_run("Spot-check Tier 1: carter.manucy@nreca.coop, venkat.banunarayanan@nreca.coop, wayne.mcgurk@nreca.coop, patti.metro@nreca.coop, meredith.miller@nreca.coop")

p = doc.add_paragraph()
run = p.add_run("Baseline reference: ")
run.bold = True
p.add_run("Section 2 -- Manucy (2 breaches), Banunarayanan (1 breach), McGurk (1 breach, email only), McNamara (1 breach)")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Personnel Affected", "Breach Name", "Data Exposed", "Priority", "Action Taken"],
    [
        ["2026-02-17", "No new breaches", "--", "--", "--", "No action required."],
        ["2026-02-24", "No new breaches", "--", "--", "--", "No action required."],
        ["", "", "", "", "", ""],
    ],
    [3, 3, 3, 3, 2, 4],
)

# --- Check 5 ---
doc.add_heading("Check 5: Alert Backlog Processing", level=3)

add_checkbox_item(doc, "Process deferred items from daily triage")

p = doc.add_paragraph()
run = p.add_run("Tool: ")
run.bold = True
p.add_run("Security Alerts email folder / Slack #security-monitoring channel")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Items in Backlog", "Items Closed", "Items Escalated"],
    [
        ["2026-02-17", "3", "2 (false positives)", "1 (FortiGate CVE -- already handled)"],
        ["2026-02-24", "1", "1 (NRECA news mention, non-security)", "0"],
        ["", "", "", ""],
    ],
    [3, 4, 5, 5],
)

# --- Check 6 ---
doc.add_heading("Check 6: Baseline Update", level=3)

add_checkbox_item(doc, "Update baseline document if any changes found")

p = doc.add_paragraph()
run = p.add_run("Time: ")
run.bold = True
p.add_run("~5 min")

add_table(doc,
    ["Date", "Section Updated", "Change Description"],
    [
        ["2026-02-17", "Vulnerabilities (M4)", "FortiGate patched to 7.4.9. CVE-2025-59718 status changed to Patched."],
        ["2026-02-24", "Attack Surface (M2)", "New subdomain newportal.cooperative.com added as Needs Investigation."],
        ["", "", ""],
    ],
    [3, 5, 10],
)

# ============================================================
# Monthly Monitoring Checklist
# ============================================================
doc.add_heading("Monthly Monitoring Checklist", level=2)

p = doc.add_paragraph()
run = p.add_run("Target time: ")
run.bold = True
p.add_run("60-90 minutes  |  ")
run2 = p.add_run("Schedule: ")
run2.bold = True
p.add_run("First Monday  |  ")
run3 = p.add_run("Owner: ")
run3.bold = True
p.add_run("[assigned analyst]")

# --- Monthly Check 1 ---
doc.add_heading("Check 1: Full Subdomain Re-enumeration (20 min)", level=3)

add_checkbox_item(doc, "Run full subdomain enumeration on all NRECA domains")

p = doc.add_paragraph()
run = p.add_run("Tools: ")
run.bold = True
p.add_run("crt.sh, DNSDumpster, SecurityTrails")

p = doc.add_paragraph()
run = p.add_run("Domains: ")
run.bold = True
p.add_run("electric.coop, cooperative.com, nreca.coop, nrecainternational.coop")

add_table(doc,
    ["Date", "New Subdomains", "Removed Subdomains", "Service Changes", "Action Taken"],
    [
        ["2026-03-03", "newportal.cooperative.com (confirmed from weekly check)", "test.community.cooperative.com (decommissioned)", "okta.cooperative.com cert renewed", "Updated baseline. Removed test.community entry."],
        ["", "", "", "", ""],
    ],
    [3, 4, 4, 4, 4],
)

# --- Monthly Check 2 ---
doc.add_heading("Check 2: Deep Shodan/Censys Scan (15 min)", level=3)

add_checkbox_item(doc, "Run broader Shodan/Censys queries")

p = doc.add_paragraph()
run = p.add_run("Queries: ")
run.bold = True
p.add_run('org:"NRECA" OR org:"National Rural Electric", ssl:"cooperative.com" OR ssl:"electric.coop"')

p = doc.add_paragraph()
run = p.add_run("Focus: ")
run.bold = True
p.add_run("New management interfaces and remote access on OT boundary. Check for FortiGate admin interfaces, new VPN endpoints, identity infrastructure changes.")

add_table(doc,
    ["Date", "New Services", "Changed Services", "Action Taken"],
    [
        ["2026-03-03", "No new services", "FortiGate now shows 7.4.9 (expected from patch)", "Confirmed patch reflected in Shodan. No unexpected changes."],
        ["", "", "", ""],
    ],
    [3, 5, 5, 5],
)

# --- Monthly Check 3 ---
doc.add_heading("Check 3: Vulnerability Re-correlation (20 min)", level=3)

add_checkbox_item(doc, "Re-run vulnerability queries for all inventoried products")

p = doc.add_paragraph()
run = p.add_run("Tools: ")
run.bold = True
p.add_run("NVD, CISA KEV, Fortinet PSIRT, ICS Advisory Project")

p = doc.add_paragraph()
run = p.add_run("Products: ")
run.bold = True
p.add_run("FortiOS 7.4.9, Okta (identity), PingFederate (identity), SharePoint (member portal)")

add_table(doc,
    ["Date", "New CVEs", "KEV Changes", "Priority Updates", "Action Taken"],
    [
        ["2026-03-03", "CVE-2026-24858 (FortiOS follow-on SSO bypass)", "Not yet in KEV", "P1 -- affects patched 7.4.9 devices", "Monitoring for KEV addition. Fortinet advisory FG-IR-26-XXX reviewed."],
        ["", "", "", "", ""],
    ],
    [3, 4, 4, 4, 4],
)

# --- Monthly Check 4 ---
doc.add_heading("Check 4: Personnel Exposure Refresh (15 min)", level=3)

add_checkbox_item(doc, "Re-check Tier 1/2 personnel against breach databases")
add_checkbox_item(doc, "Review personnel changes")

add_table(doc,
    ["Date", "New Breaches", "Personnel Changes", "Action Taken"],
    [
        ["2026-03-03", "No new breaches for Tier 1/2", "No personnel changes identified", "No action required."],
        ["", "", "", ""],
    ],
    [3, 5, 5, 5],
)

# --- Monthly Check 5 ---
doc.add_heading("Check 5: Alert Configuration Review (10 min)", level=3)

add_checkbox_item(doc, "Verify all push alerts are active and delivering")
add_checkbox_item(doc, "Review query relevance")

add_table(doc,
    ["Date", "Alerts Verified", "Changes Made"],
    [
        ["2026-03-03", "5 Google Alerts active, CISA ICS subscription active, Fortinet PSIRT active", "Added Google Alert for CVE-2026-24858 follow-on vulnerability."],
        ["", "", ""],
    ],
    [3, 6, 8],
)

# ============================================================
# Cycle Summary
# ============================================================
doc.add_heading("Cycle Summary", level=2)

add_table(doc,
    ["Cycle Date", "Type", "Total Findings", "P0/P1", "Baseline Updates", "Time Spent", "Analyst"],
    [
        ["2026-02-17", "Weekly", "2", "1 (P0)", "2", "25 min", "[analyst]"],
        ["2026-02-24", "Weekly", "1", "0", "1", "20 min", "[analyst]"],
        ["2026-03-03", "Monthly", "3", "1 (P1)", "3", "75 min", "[analyst]"],
        ["", "", "", "", "", "", ""],
    ],
    [3, 2, 2.5, 2, 2.5, 2, 3],
)

doc.save("/home/cutaway/Projects/ot-osint-program-workshop/examples/downloads/nreca-monitoring-checklist.docx")
print("nreca-monitoring-checklist.docx created")
