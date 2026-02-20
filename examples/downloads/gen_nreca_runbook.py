"""Generate nreca-runbook.docx pre-filled with NRECA example data."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)

for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    hs.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A365D", qn("w:val"): "clear",
        })
        shading.append(shading_elm)
    for r, row_data in enumerate(rows, 1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_step(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_box = p.add_run("\u2610 ")
    run_box.font.size = Pt(11)
    run_label = p.add_run(f"[{label}] ")
    run_label.font.bold = True
    run_label.font.size = Pt(11)
    if label == "CORE":
        run_label.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    else:
        run_label.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    return p


def add_detail(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(value)
    run2.font.size = Pt(10)
    return p


# ============================================================
# Title
# ============================================================
doc.add_heading("Operational Runbook -- NRECA Example", level=1)

p = doc.add_paragraph()
p.add_run("NOTE: ").bold = True
p.add_run("This is a worked example using NRECA (National Rural Electric Cooperative Association) as the target organization. All data is derived from publicly available sources. Breach results are hypothetical examples created for the workshop.")

doc.add_paragraph()

p = doc.add_paragraph()
run_core = p.add_run("[CORE]")
run_core.bold = True
run_core.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
p.add_run(" -- Runs no matter what. Part of the minimum viable monitoring set.")

p = doc.add_paragraph()
run_full = p.add_run("[FULL]")
run_full.bold = True
run_full.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
p.add_run(" -- Runs during normal operations. Can be deferred when time is constrained.")

doc.add_paragraph()

# ============================================================
# Program Information
# ============================================================
doc.add_heading("Program Information", level=2)

info_fields = [
    ("Organization", "NRECA (National Rural Electric Cooperative Association)"),
    ("Program owner", "Carter Manucy, Director of Cybersecurity"),
    ("Backup owner", "Adrian McNamara, Cybersecurity Program Manager"),
    ("Escalation contact (P0/P1)", "Carter Manucy, [phone], carter.manucy@nreca.coop"),
    ("Baseline document location", "[internal SharePoint / document management path]"),
    ("Last runbook review", "[date]"),
]
for label, value in info_fields:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

doc.add_paragraph()

# ============================================================
# Daily Alert Triage
# ============================================================
doc.add_heading("Daily Alert Triage (5-10 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("Adrian McNamara  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("Carter Manucy")

add_step(doc, "CORE", "Open Security Alerts email folder (filtered for Google Alerts sender)")

add_step(doc, "CORE", "Review new Google Alerts for these 5 configured queries:")
add_detail(doc, "Query 1", '"NRECA" breach OR hack OR vulnerability OR "data leak" OR ransomware')
add_detail(doc, "Query 2", '"electric cooperative" SCADA OR ICS OR "control system" vulnerability OR breach OR attack')
add_detail(doc, "Query 3", "nreca.coop OR electric.coop OR cooperative.com exposed OR leak OR breach OR credential")
add_detail(doc, "Query 4", 'FortiGate OR FortiOS vulnerability OR "zero day" OR exploit OR CVE')
add_detail(doc, "Query 5", '"nreca.coop" OR "electric.coop" site:pastebin.com OR site:github.com OR site:reddit.com')

add_step(doc, "CORE", "Review CISA ICS Advisory emails -- check for advisories affecting Fortinet, Okta, PingFederate")

add_step(doc, "CORE", "Triage each alert:")
add_detail(doc, "Known baseline item", "Mark as reviewed, no action")
add_detail(doc, "New finding, low priority", "Add to weekly review queue with one-line note")
add_detail(doc, "New finding, elevated priority", "Investigate immediately using P0-P3 framework")
add_detail(doc, "New finding, critical (P0)", "Escalate to Carter Manucy immediately")

add_step(doc, "CORE", "Log triage decisions: date, source, summary, disposition")

doc.add_paragraph()

# ============================================================
# Weekly Active Monitoring
# ============================================================
doc.add_heading("Weekly Active Monitoring (25-30 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("Adrian McNamara  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("Carter Manucy  |  ")
run3 = p.add_run("Day: ")
run3.bold = True
p.add_run("Monday")

p = doc.add_paragraph("Checks ordered by OT relevance:")
p.italic = True

add_step(doc, "CORE", "Shodan/Censys remote access check (5 min)")
add_detail(doc, "Query", 'Shodan: product:"FortiGate" port:443 "FortiOS" net:[NRECA IP ranges]')
add_detail(doc, "Also check", "Censys: services.software.product=FortiOS AND ip:[NRECA IP ranges]")
add_detail(doc, "Compare to", "Baseline Section 1 -- FortiGate SSL-VPN on port 443, currently FortiOS 7.4.9")
add_detail(doc, "Escalate if", "New ports, service changes, FortiOS version change, or unexpected admin interface")

add_step(doc, "CORE", "CISA KEV review (5 min)")
add_detail(doc, "Query", "Review CISA KEV catalog additions from past 7 days")
add_detail(doc, "Cross-reference", "FortiOS, Okta, PingFederate, SharePoint, any cooperative sector products")
add_detail(doc, "Escalate if", "Any KEV match on an internet-exposed NRECA asset (P0)")

add_step(doc, "FULL", "Certificate transparency check (5 min)")
add_detail(doc, "Queries", "crt.sh: %.electric.coop, %.cooperative.com, %.nreca.coop, %.nrecainternational.coop")
add_detail(doc, "Compare to", "Baseline Section 1 -- electric.coop (5 subdomains), cooperative.com (100+ subdomains)")
add_detail(doc, "Escalate if", "New subdomains suggesting remote access (vpn*, ras*, remote*) or identity infrastructure")

add_step(doc, "FULL", "Breach database check (5 min)")
add_detail(doc, "Query", "HIBP spot-check Tier 1: carter.manucy@nreca.coop, venkat.banunarayanan@nreca.coop, wayne.mcgurk@nreca.coop, patti.metro@nreca.coop, meredith.miller@nreca.coop")
add_detail(doc, "Compare to", "Baseline Section 2 -- Manucy (2 breaches), Banunarayanan (1), McGurk (1, email only), McNamara (1)")
add_detail(doc, "Escalate if", "Tier 1 personnel in new breach with password/hash exposure")

add_step(doc, "FULL", "Alert backlog processing (5 min)")
add_detail(doc, "Action", "Review items deferred during daily triage from Security Alerts folder")
add_detail(doc, "Rule", "No item remains in backlog longer than one week")

add_step(doc, "FULL", "Baseline update (5 min)")
add_detail(doc, "Action", "Update baseline document for any changes discovered this week")
add_detail(doc, "Log", "Record date and nature of each change in the Change Log sheet")

doc.add_paragraph()

# ============================================================
# Monthly Full Review
# ============================================================
doc.add_heading("Monthly Full Review (60-90 minutes)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("Carter Manucy  |  ")
run2 = p.add_run("Backup: ")
run2.bold = True
p.add_run("Adrian McNamara  |  ")
run3 = p.add_run("Week: ")
run3.bold = True
p.add_run("First Monday")

add_step(doc, "FULL", "Attack surface re-enumeration (20 min)")
add_detail(doc, "Tools", "crt.sh, DNSDumpster, SecurityTrails")
add_detail(doc, "Domains", "electric.coop, cooperative.com, nreca.coop, nrecainternational.coop")
add_detail(doc, "Focus", "New subdomains under cooperative.com (primary attack surface with 100+ subdomains)")
add_detail(doc, "Action", "Compare complete results to baseline. Investigate new subdomains. Remove decommissioned entries")

add_step(doc, "FULL", "Deep Shodan/Censys scan (15 min)")
add_detail(doc, "Queries", 'org:"NRECA" OR org:"National Rural Electric", ssl:"cooperative.com" OR ssl:"electric.coop"')
add_detail(doc, "Focus", "New management interfaces, identity infrastructure changes (Okta, PingFederate), remote access endpoints")
add_detail(doc, "Check", "Blue/green deployment pattern -- monitor both blue.* and production subdomains for drift")

add_step(doc, "CORE", "Vulnerability re-correlation (20 min)")
add_detail(doc, "Products", "FortiOS 7.4.9, Okta (identity), PingFederate (identity), SharePoint (member portal)")
add_detail(doc, "Sources", "NVD, CISA KEV, Fortinet PSIRT (fortiguard.fortinet.com/psirt), ICS Advisory Project")
add_detail(doc, "Action", "Update vulnerability correlation table. Re-evaluate all priority ratings")
add_detail(doc, "Watch for", "Follow-on CVEs like CVE-2026-24858 that affect patched versions")

add_step(doc, "FULL", "Personnel exposure refresh (15 min)")
add_detail(doc, "Re-check", "All 5 Tier 1 (Manucy, Banunarayanan, Metro, Miller, McGurk) and 3 Tier 2 (Singh, McNamara, Loving) against HIBP")
add_detail(doc, "Review", "cooperative.com and electric.coop for new OT-relevant staff postings")
add_detail(doc, "Team inboxes", "Verify membersecurity@nreca.coop and DistributionAutomation@nreca.coop are not in new breaches")

add_step(doc, "FULL", "Alert configuration review (10 min)")
add_detail(doc, "Verify", "All 5 Google Alert queries still active and delivering")
add_detail(doc, "Verify", "CISA ICS Advisory subscription, Fortinet PSIRT subscription active")
add_detail(doc, "Update", "Add queries for any new products, domains, or vendors discovered this month")

doc.add_paragraph()

# ============================================================
# Quarterly Program Assessment
# ============================================================
doc.add_heading("Quarterly Program Assessment (2-3 hours)", level=2)

p = doc.add_paragraph()
run = p.add_run("Owner: ")
run.bold = True
p.add_run("Carter Manucy  |  ")
run2 = p.add_run("Quarter: ")
run2.bold = True
p.add_run("[Q1/Q2/Q3/Q4]")

add_step(doc, "FULL", "Full baseline refresh (60 min)")
add_detail(doc, "Action", "Complete re-run of Modules 2-4 discovery as if starting fresh")
add_detail(doc, "Compare", "Comprehensive results to current baseline across all 4 NRECA domains")
add_detail(doc, "Reclassify", "All entries using the four-category system (Known-Good, Accepted Risk, Needs Remediation, Needs Investigation)")

add_step(doc, "FULL", "Program metrics (30 min)")
add_detail(doc, "Total alerts received vs. actionable findings", "___")
add_detail(doc, "Average triage response time", "___")
add_detail(doc, "Baseline changes this quarter", "___")
add_detail(doc, "Open remediation items and aging", "___")
add_detail(doc, "FortiGate-specific", "Time from CVE disclosure to patch applied: ___")

add_step(doc, "FULL", "Stakeholder briefing (30 min)")
add_detail(doc, "Audience", "NRECA cybersecurity leadership, CIO")
add_detail(doc, "Content", "Key findings, remediation progress, emerging risks to cooperative network, resource requests")

add_step(doc, "FULL", "Runbook review (30 min)")
add_detail(doc, "Check", "Are queries still returning relevant results?")
add_detail(doc, "Check", "Have any NRECA domains or key personnel changed?")
add_detail(doc, "Check", "Are there new products in the network that need vulnerability monitoring?")
add_detail(doc, "Action", "Update cadences, queries, contacts, and product inventory as needed")
add_detail(doc, "Action", "Update this runbook's 'Last runbook review' date")

doc.add_paragraph()

# ============================================================
# Change Log
# ============================================================
doc.add_heading("Change Log", level=2)

add_table(doc,
    ["Date", "Change", "Changed By"],
    [
        ["[today]", "Initial runbook created from workshop", "[analyst]"],
        ["", "", ""],
        ["", "", ""],
    ],
    [3, 10, 4],
)

doc.save("/home/cutaway/Projects/ot-osint-program-workshop/examples/downloads/nreca-runbook.docx")
print("nreca-runbook.docx created")
